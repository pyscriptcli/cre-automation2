import os
import time
import subprocess
import tempfile
import streamlit as st
import requests
import json
import pandas as pd
import datetime
import re
from io import BytesIO
import PyPDF2
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
import streamlit.components.v1 as components

# ========== CONFIG ==========
st.set_page_config(page_title="Project Echo", layout="wide", initial_sidebar_state="collapsed")[cite: 1]

# --- PROGRAMMATIC LIGHT MODE & 200MB LIMIT ---
_config_dir = ".streamlit"[cite: 1]
_config_file = os.path.join(_config_dir, "config.toml")[cite: 1]
os.makedirs(_config_dir, exist_ok=True)[cite: 1]
with open(_config_file, "w", encoding="utf-8") as f:[cite: 1]
    f.write('[theme]\nbase="light"\n[server]\nmaxUploadSize = 200\n')[cite: 1]

# Simple Password Gate at the start of app execution (Read strictly from Streamlit Cloud Secrets)
if "authenticated" not in st.session_state:[cite: 1]
    st.session_state["authenticated"] = False[cite: 1]

if not st.session_state["authenticated"]:[cite: 1]
    pw_input = st.text_input("Enter Team Access Key to use Project Echo:", type="password")[cite: 1]
    if st.button("Log In"):[cite: 1]
        configured_password = str(st.secrets.get("APP_PASSWORD", "crd3ch0")).strip()
        if pw_input == configured_password:[cite: 1]
            st.session_state["authenticated"] = True[cite: 1]
            st.rerun()[cite: 1]
        else:
            st.error("Invalid access key. Contact the administrator.")[cite: 1]
    st.stop()  # Prevents unauthorized access to app functions[cite: 1]

# API Keys loaded strictly from Streamlit Cloud Secrets
DEEPSEEK_API_KEY = str(st.secrets.get("DEEPSEEK_API_KEY", "")).strip()
DEEPSEEK_CHAT_URL = "https://api.deepseek.com/chat/completions"[cite: 1]

GROQ_API_KEY = str(st.secrets.get("GROQ_API_KEY", "")).strip()
GROQ_AUDIO_URL = "https://api.groq.com/openai/v1/audio/transcriptions"[cite: 1]

OPENAI_API_KEY = str(st.secrets.get("OPENAI_API_KEY", "")).strip()
OPENAI_AUDIO_URL = "https://api.openai.com/v1/audio/transcriptions"[cite: 1]

CRD_MEMBERS = [[cite: 1]
    "Sondi Tuazon",[cite: 1]
    "Kristina Balajadia",[cite: 1]
    "Meliza Zapata",[cite: 1]
    "Dykstra Pineda",[cite: 1]
    "Cedtrix Rena",[cite: 1]
    "Carlo Medina",[cite: 1]
    "Dave Policarpio",[cite: 1]
    "Irish Rima"[cite: 1]
][cite: 1]

LOCATION_PRESETS = [[cite: 1]
    "— Select a Preset (Optional) —",[cite: 1]
    "GreatWork Mega Tower 32F - Secret Room",[cite: 1]
    "GreatWork Mega Tower 32F - Small Meeting Room",[cite: 1]
    "GreatWork Mega Tower 24F - Meeting Room",[cite: 1]
    "GreatWork Mega Tower 32F - Board Room",[cite: 1]
    "GreatWork Mega Tower 32F - Co-working",[cite: 1]
    "Online Meeting"[cite: 1]
][cite: 1]

# Initialize Session State Variables
if "transcript" not in st.session_state:[cite: 1]
    st.session_state["transcript"] = ""[cite: 1]
if "df" not in st.session_state:[cite: 1]
    st.session_state["df"] = pd.DataFrame(columns=["Discussion Points", "Action Plan", "Indicative Delivery Date", "Person-in-charge"])[cite: 1]
if "other_discussions" not in st.session_state:[cite: 1]
    st.session_state["other_discussions"] = ""[cite: 1]
if "show_settings" not in st.session_state:[cite: 1]
    st.session_state["show_settings"] = False[cite: 1]
if "tokens_used" not in st.session_state:[cite: 1]
    st.session_state["tokens_used"] = 0[cite: 1]
if "last_api_call" not in st.session_state:[cite: 1]
    st.session_state["last_api_call"] = None[cite: 1]
if "selected_engine" not in st.session_state:[cite: 1]
    st.session_state["selected_engine"] = "AI - DeepSeek"[cite: 1]
if "chat_history" not in st.session_state:[cite: 1]
    st.session_state["chat_history"] = [][cite: 1]

# ========== CUSTOM CSS ==========
CUSTOM_CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Montserrat:wght@400;500;600&family=Playfair+Display:ital,wght@1,400;1,500;1,600&display=swap');

html, body, [class*="css"] {
    font-family: 'Montserrat', sans-serif !important;
}

.stApp {
    background-color: #F4F2EC; 
    background-image: 
        linear-gradient(rgba(0, 0, 0, 0.02) 1px, transparent 1px),
        linear-gradient(90deg, rgba(0, 0, 0, 0.02) 1px, transparent 1px);
    background-size: 80px 80px;
    color: #333333;
}

.stApp > header { display: none !important; }
.block-container { padding-top: 5.5rem !important; }

/* Fixed Topbar */
.echo-topbar-wrapper {
    position: fixed; top: 0; left: 0; right: 0; height: 60px;
    background-color: #161616;
    border-bottom: 1px solid #333333;
    z-index: 999990; box-shadow: 0 4px 15px rgba(0,0,0,0.2);
    display: flex; align-items: center; justify-content: flex-start;
    padding: 0 2rem;
}

.echo-title {
    font-family: 'Playfair Display', serif !important;
    font-style: italic !important; font-weight: 400 !important;
    font-size: 1.35rem !important; color: #FFFFFF !important; margin: 0 !important;
}
.echo-title span { color: #D4AF37 !important; }

h3 {
    font-family: 'Playfair Display', serif !important;
    font-style: italic !important; font-weight: 400 !important; 
    color: #1A2B4C !important; letter-spacing: 0.02em; margin-bottom: 0.25rem; font-size: 1.25rem !important;
}

/* Playfair Display Styling for Labels & Tab Headers */
.playfair-label {
    font-family: 'Playfair Display', serif !important;
    font-style: italic !important;
    font-weight: 400 !important;
    color: #1A2B4C !important;
    font-size: 1.05rem !important;
    margin-bottom: 0.25rem !important;
    display: block;
}

/* Target Tab Text & Make it Playfair */
button[data-baseweb="tab"] p {
    font-family: 'Playfair Display', serif !important;
    font-style: italic !important;
    font-weight: 400 !important;
    color: #1A2B4C !important;
    font-size: 1.05rem !important;
}

[data-testid="stVerticalBlockBorderWrapper"] {
    background-color: #FFFFFF !important; border-radius: 12px !important;
    box-shadow: 0 8px 30px rgba(0, 0, 0, 0.04) !important;
    border: 1px solid rgba(0, 0, 0, 0.04) !important; 
    padding: 1.25rem !important; margin-bottom: 1rem !important;
}

/* Uniform Small Pill Buttons */
.stButton > button, .stDownloadButton > button {
    background-color: #222222 !important; 
    color: #FFFFFF !important;
    border: 1px solid #444444 !important; 
    border-radius: 50px !important; 
    font-family: 'Montserrat', sans-serif !important; 
    font-weight: 500 !important;
    font-size: 0.82rem !important;
    letter-spacing: 0.5px; 
    padding: 0.35rem 1.1rem !important;
    min-height: 34px !important;
    height: 34px !important;
    transition: all 0.25s ease !important; 
    width: 100% !important;
}

.stButton > button:hover, .stDownloadButton > button:hover {
    border-color: #D4AF37 !important; 
    color: #D4AF37 !important;
    background-color: #1A1A1A !important;
}

button[key^="del_"] {
    min-height: 32px !important;
    height: 32px !important;
    padding: 0.2rem 0.5rem !important;
    background-color: #F7F5F0 !important;
    color: #A03030 !important;
    border: 1px solid rgba(160, 48, 48, 0.25) !important;
}

button[key^="del_"]:hover {
    background-color: #A03030 !important;
    color: #FFFFFF !important;
    border-color: #A03030 !important;
}

div[data-testid="stButton"]:has(button[key="card_settings_btn"]) {
    display: flex !important;
    justify-content: flex-end !important;
}

button[key="card_settings_btn"] {
    background-color: transparent !important;
    border: 1px solid #C5A059 !important;
    border-radius: 50% !important;
    width: 32px !important;
    height: 32px !important;
    min-height: 32px !important;
    padding: 0 !important;
    margin: 0 !important;
    display: flex !important;
    align-items: center !important;
    justify-content: center !important;
    cursor: pointer !important;
    box-shadow: none !important;
}

button[key="card_settings_btn"]::before {
    content: "";
    display: inline-block;
    width: 17px;
    height: 17px;
    background-color: #C5A059;
    -webkit-mask: url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='currentColor' stroke-width='2' stroke-linecap='round' stroke-linejoin='round'%3E%3Ccircle cx='12' cy='12' r='3'%3E%3C/circle%3E%3Cpath d='M19.4 15a1.65 1.65 0 0 0 .33 1.82l.06.06a2 2 0 0 1 0 2.83 2 2 0 0 1-2.83 0l-.06-.06a1.65 1.65 0 0 0-1.82-.33 1.65 1.65 0 0 0-1 1.51V21a2 2 0 0 1-2 2 2 2 0 0 1-2-2v-.09A1.65 1.65 0 0 0 9 19.4a1.65 1.65 0 0 0-1.82.33l-.06.06a2 2 0 0 1-2.83 0 2 2 0 0 1 0-2.83l.06-.06a1.65 1.65 0 0 0 .33-1.82 1.65 1.65 0 0 0-1.51-1H3a2 2 0 0 1-2-2 2 2 0 0 1 2-2h.09A1.65 1.65 0 0 0 4.6 9a1.65 1.65 0 0 0-.33-1.82l-.06-.06a2 2 0 0 1 0-2.83 2 2 0 0 1 2.83 0l.06.06a1.65 1.65 0 0 0 1.82.33H9a1.65 1.65 0 0 0 1-1.51V3a2 2 0 0 1 2-2 2 2 0 0 1 2 2v.09a1.65 1.65 0 0 0 1 1.51 1.65 1.65 0 0 0 1.82-.33l.06-.06a2 2 0 0 1 2.83 0 2 2 0 0 1 0 2.83l-.06.06a1.65 1.65 0 0 0-.33 1.82V9a1.65 1.65 0 0 0 1.51 1H21a2 2 0 0 1 2 2 2 2 0 0 1-2 2h-.09a1.65 1.65 0 0 0-1.51 1z'%3E%3C/path%3E%3C/svg%3E") no-repeat center;
    mask: url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='currentColor' stroke-width='2' stroke-linecap='round' stroke-linejoin='round'%3E%3Ccircle cx='12' cy='12' r='3'%3E%3C/circle%3E%3Cpath d='M19.4 15a1.65 1.65 0 0 0 .33 1.82l.06.06a2 2 0 0 1 0 2.83 2 2 0 0 1-2.83 0l-.06-.06a1.65 1.65 0 0 0-1.82-.33 1.65 1.65 0 0 0-1 1.51V21a2 2 0 0 1-2 2 2 2 0 0 1-2-2v-.09A1.65 1.65 0 0 0 9 19.4a1.65 1.65 0 0 0-1.82.33l-.06.06a2 2 0 0 1 2.83 0 2 2 0 0 1 0 2.83l-.06.06a1.65 1.65 0 0 0-.33 1.82V9a1.65 1.65 0 0 0 1.51 1H21a2 2 0 0 1 2 2 2 2 0 0 1-2 2h-.09a1.65 1.65 0 0 0-1.51 1z'%3E%3C/path%3E%3C/svg%3E") no-repeat center;
    -webkit-mask-size: contain;
    mask-size: contain;
    transition: background-color 0.2s ease;
}

/* Auto-wrapping text inputs */
.stTextArea textarea {
    font-size: 0.92rem !important;[cite: 1]
    line-height: 1.45 !important;[cite: 1]
    border-radius: 8px !important;[cite: 1]
}

/* Time Picker Clean Layout */
[data-testid="column"] .stSelectbox {
    margin-bottom: 0 !important;[cite: 1]
}

/* Sophisticated, Mature & Minimalist Ask Echo Chat UI */
.chat-container {
    display: flex;
    flex-direction: column;
    gap: 0.85rem;
    margin-top: 0.5rem;
    max-height: 380px;
    overflow-y: auto;
    padding-right: 0.35rem;
}

.chat-bubble-ai-wrap {
    display: flex;
    justify-content: flex-start;
    align-items: flex-start;
    gap: 0.5rem;
}

.chat-bubble-ai {
    background-color: #FAF9F5;
    color: #222222;
    padding: 0.85rem 1.1rem;
    border-radius: 4px 14px 14px 14px;
    max-width: 88%;
    font-size: 0.88rem;
    line-height: 1.55;
    border: 1px solid rgba(212, 175, 55, 0.35);
    box-shadow: 0 2px 8px rgba(0,0,0,0.02);
}

.chat-bubble-ai .ai-header {
    font-size: 0.72rem;
    text-transform: uppercase;
    letter-spacing: 0.8px;
    font-weight: 600;
    color: #A07828;
    margin-bottom: 0.35rem;
}

.chat-bubble-user-wrap {
    display: flex;
    justify-content: flex-end;
}

.chat-bubble-user {
    background-color: #ECE8DD;
    color: #1A1A1A;
    padding: 0.65rem 1rem;
    border-radius: 14px 14px 4px 14px;
    max-width: 85%;
    font-size: 0.88rem;
    line-height: 1.45;
    border: 1px solid rgba(0,0,0,0.06);
}
</style>
"""

# ========== SVG ICONS ==========
SVG_ALERT = """<svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="#D4AF37" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M10.29 3.86L1.82 18a2 2 0 0 0 1.71 3h16.94a2 2 0 0 0 1.71-3L13.71 3.86a2 2 0 0 0-3.42 0z"></path><line x1="12" y1="9" x2="12" y2="13"></line><line x1="12" y1="17" x2="12.01" y2="17"></line></svg>"""[cite: 1]

# ========== CORE LOGIC ==========
def extract_text_from_file(uploaded_file):
    try:
        if uploaded_file.name.endswith('.txt'):[cite: 1]
            return uploaded_file.getvalue().decode("utf-8")[cite: 1]
        elif uploaded_file.name.endswith('.pdf'):[cite: 1]
            reader = PyPDF2.PdfReader(uploaded_file)[cite: 1]
            text = ""[cite: 1]
            for page in reader.pages:[cite: 1]
                text += page.extract_text() + "\n"[cite: 1]
            return text[cite: 1]
        elif uploaded_file.name.endswith('.docx'):[cite: 1]
            doc = Document(uploaded_file)[cite: 1]
            return "\n".join([para.text for para in doc.paragraphs])[cite: 1]
        return ""[cite: 1]
    except Exception as e:[cite: 1]
        st.error(f"Error reading file: {e}")[cite: 1]
        return ""[cite: 1]

def _call_openai_transcribe(audio_bytes, filename="audio.mp3"):
    if not OPENAI_API_KEY:[cite: 1]
        st.error("OpenAI API Key is missing. Please add it to your Streamlit Cloud Secrets.")[cite: 1]
        return None[cite: 1]
    headers = {"Authorization": f"Bearer {OPENAI_API_KEY}"}[cite: 1]
    files = {"file": (filename, audio_bytes), "model": (None, "gpt-4o-mini-transcribe"), "response_format": (None, "json")}[cite: 1]
    try:
        resp = requests.post(OPENAI_AUDIO_URL, headers=headers, files=files, timeout=180)[cite: 1]
        if resp.status_code == 200:[cite: 1]
            return resp.json().get("text", "")[cite: 1]
        st.error(f"OpenAI fallback error: {resp.text}")[cite: 1]
        return None[cite: 1]
    except Exception as e:[cite: 1]
        st.error(f"OpenAI connection error: {e}")[cite: 1]
        return None[cite: 1]

def _call_groq_whisper(audio_bytes, filename="audio.mp3"):
    headers = {"Authorization": f"Bearer {GROQ_API_KEY}"}[cite: 1]
    files = {"file": (filename, audio_bytes), "model": (None, "whisper-large-v3-turbo"), "response_format": (None, "json")}[cite: 1]
    try:
        resp = requests.post(GROQ_AUDIO_URL, headers=headers, files=files, timeout=60)[cite: 1]
        if resp.status_code == 200:[cite: 1]
            return resp.json().get("text", "")[cite: 1]
        return None[cite: 1]
    except:
        return None[cite: 1]

def transcribe_audio_pipeline(audio_bytes, original_filename, progress_bar, status_placeholder):
    progress_bar.progress(10, text="Preprocessing audio container (10%)...")[cite: 1]
    
    ext = os.path.splitext(original_filename)[1] or ".m4a"[cite: 1]
    with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as src:[cite: 1]
        src.write(audio_bytes)[cite: 1]
        src_path = src.name[cite: 1]

    compressed_mp3 = src_path + "_compressed.mp3"[cite: 1]
    progress_bar.progress(25, text="Compressing audio to 16kHz Mono 24k MP3 (25%)...")[cite: 1]

    try:
        cmd = [[cite: 1]
            "ffmpeg", "-y",[cite: 1]
            "-threads", "1",[cite: 1]
            "-i", src_path,[cite: 1]
            "-vn",[cite: 1]
            "-ac", "1",[cite: 1]
            "-ar", "16000",[cite: 1]
            "-c:a", "libmp3lame",[cite: 1]
            "-b:a", "24k",[cite: 1]
            compressed_mp3[cite: 1]
        ][cite: 1]
        res = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)[cite: 1]
        if res.returncode != 0:[cite: 1]
            st.error(f"FFmpeg compression error: {res.stderr[:200]}")[cite: 1]
            return None[cite: 1]

        comp_size_mb = os.path.getsize(compressed_mp3) / (1024 * 1024)[cite: 1]
        progress_bar.progress(45, text="Evaluating audio duration & routing (45%)...")[cite: 1]

        if comp_size_mb <= 10.0 and GROQ_API_KEY:[cite: 1]
            status_placeholder.info("Processing via Groq Whisper Primary...")[cite: 1]
            progress_bar.progress(70, text="Transcribing via Groq Whisper (70%)...")[cite: 1]
            with open(compressed_mp3, "rb") as f:[cite: 1]
                c_bytes = f.read()[cite: 1]
            text = _call_groq_whisper(c_bytes, "audio.mp3")[cite: 1]
            if text:[cite: 1]
                progress_bar.progress(100, text="Transcription completed (100%)!")[cite: 1]
                status_placeholder.empty()[cite: 1]
                return text[cite: 1]
            status_placeholder.warning("Groq rate limit reached. Switching automatically to OpenAI...")[cite: 1]

        status_placeholder.info("Processing recording via OpenAI...")[cite: 1]
        progress_bar.progress(55, text="Preparing audio segments for OpenAI (55%)...")[cite: 1]
        
        segment_pattern = src_path + "_seg_%03d.mp3"[cite: 1]
        subprocess.run([[cite: 1]
            "ffmpeg", "-y", "-i", compressed_mp3,[cite: 1]
            "-f", "segment", "-segment_time", "600", "-c", "copy",[cite: 1]
            segment_pattern[cite: 1]
        ], stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)[cite: 1]

        seg_dir = os.path.dirname(src_path)[cite: 1]
        base_name = os.path.basename(src_path) + "_seg_"[cite: 1]
        segments = sorted([os.path.join(seg_dir, f) for f in os.listdir(seg_dir) if f.startswith(base_name)])[cite: 1]

        full_transcript = [][cite: 1]
        total_segs = len(segments)[cite: 1]

        for idx, seg in enumerate(segments):[cite: 1]
            pct = int(55 + ((idx + 1) / total_segs) * 40)[cite: 1]
            progress_bar.progress(pct, text=f"Transcribing segment {idx + 1} of {total_segs} ({pct}%)...")[cite: 1]
            
            with open(seg, "rb") as f:[cite: 1]
                seg_bytes = f.read()[cite: 1]
            t = _call_openai_transcribe(seg_bytes, f"part_{idx}.mp3")[cite: 1]
            if t:[cite: 1]
                full_transcript.append(t)[cite: 1]
            time.sleep(0.2)[cite: 1]
            try:
                os.remove(seg)[cite: 1]
            except:
                pass[cite: 1]

        progress_bar.progress(100, text="Transcription completed successfully (100%)!")[cite: 1]
        time.sleep(0.3)[cite: 1]
        status_placeholder.empty()[cite: 1]
        return " ".join(full_transcript)[cite: 1]

    except Exception as e:[cite: 1]
        st.error(f"Audio processing failure: {e}")[cite: 1]
        return None[cite: 1]
    finally:
        if os.path.exists(src_path):[cite: 1]
            try:
                os.remove(src_path)[cite: 1]
            except:
                pass[cite: 1]
        if os.path.exists(compressed_mp3):[cite: 1]
            try:
                os.remove(compressed_mp3)[cite: 1]
            except:
                pass[cite: 1]

def normalize_llm_json_to_df(data):
    items = None[cite: 1]
    other_disc = ""[cite: 1]
    
    if isinstance(data, list):[cite: 1]
        items = data[cite: 1]
    elif isinstance(data, dict):[cite: 1]
        for key in ["table_items", "items", "minutes", "table", "data", "discussion_items", "discussions", "action_items"]:[cite: 1]
            if key in data and isinstance(data[key], list) and len(data[key]) > 0:[cite: 1]
                items = data[key][cite: 1]
                break[cite: 1]
        if items is None:[cite: 1]
            for v in data.values():[cite: 1]
                if isinstance(v, list) and len(v) > 0 and isinstance(v[0], dict):[cite: 1]
                    items = v[cite: 1]
                    break[cite: 1]
            if items is None:[cite: 1]
                items = [data][cite: 1]
                
        other_disc = str(data.get("other_discussions", "") or data.get("notes", "") or data.get("summary", ""))[cite: 1]

    if not items or not isinstance(items, list):[cite: 1]
        return None, ""[cite: 1]

    df = pd.DataFrame(items)[cite: 1]
    col_mapping = {}[cite: 1]
    for c in df.columns:[cite: 1]
        c_clean = str(c).lower().replace("_", " ").replace("-", " ")[cite: 1]
        if any(k in c_clean for k in ["discuss", "point", "topic", "milestone"]):[cite: 1]
            col_mapping[c] = "Discussion Points"[cite: 1]
        elif any(k in c_clean for k in ["action", "plan", "step", "deliverable"]):[cite: 1]
            col_mapping[c] = "Action Plan"[cite: 1]
        elif any(k in c_clean for k in ["date", "time", "delivery", "deadline"]):[cite: 1]
            col_mapping[c] = "Indicative Delivery Date"[cite: 1]
        elif any(k in c_clean for k in ["person", "charge", "pic", "assign", "who", "responsible"]):[cite: 1]
            col_mapping[c] = "Person-in-charge"[cite: 1]

    df = df.rename(columns=col_mapping)[cite: 1]
    for col in ["Discussion Points", "Action Plan", "Indicative Delivery Date", "Person-in-charge"]:[cite: 1]
        if col not in df.columns:[cite: 1]
            df[col] = ""[cite: 1]
            
    df = df[["Discussion Points", "Action Plan", "Indicative Delivery Date", "Person-in-charge"]].drop_duplicates()[cite: 1]
    return df, other_disc[cite: 1]

def extract_with_deepseek(transcript):
    if not DEEPSEEK_API_KEY:[cite: 1]
        st.error("DeepSeek API Key is missing. Please add it to your Streamlit Cloud Secrets.")[cite: 1]
        return None, ""[cite: 1]

    headers = {[cite: 1]
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}",[cite: 1]
        "Content-Type": "application/json"[cite: 1]
    }[cite: 1]

    system_prompt = ([cite: 1]
        "You are an expert executive assistant for PRIME Philippines tasked with producing comprehensive, "[cite: 1]
        "high-level executive Minutes of the Meeting (MOM). "[cite: 1]
        "The transcript contains Tagalog, English, and Taglish dialogue. "[cite: 1]
        "Analyze the full conversation context and translate all colloquial, informal, and mixed-language statements "[cite: 1]
        "into polished, high-level corporate English. "[cite: 1]
        "Synthesize all key agreements, status reports, core discussion points, definitive action plans, "[cite: 1]
        "indicative delivery timelines, and assigned persons-in-charge without omitting critical business context. "[cite: 1]
        "Output valid JSON only matching the exact schema provided."[cite: 1]
    )[cite: 1]

    user_prompt = f"""Synthesize the following meeting transcript into formal, high-level Minutes of Meeting (MOM) formatted as valid JSON:

Schema:
{{
  "table_items": [
    {{
      "Discussion Points": "Formal summary of key milestones, operational updates, or strategic topics discussed",
      "Action Plan": "Concrete, actionable executive deliverables and next steps (state 'None' if purely informational)",
      "Indicative Delivery Date": "Specific date, timeline, or 'TBD'",
      "Person-in-charge": "Designated individual, department (e.g., PRIME Philippines, Client name), or 'Unassigned'"
    }}
  ],
  "other_discussions": "High-level summary of peripheral discussions, informal remarks, or general alignment"
}}

Transcript:
{transcript[:28000]}"""[cite: 1]

    payload = {[cite: 1]
        "model": "deepseek-chat",[cite: 1]
        "messages": [[cite: 1]
            {"role": "system", "content": system_prompt},[cite: 1]
            {"role": "user", "content": user_prompt}[cite: 1]
        ],[cite: 1]
        "response_format": {"type": "json_object"},[cite: 1]
        "temperature": 0.1,[cite: 1]
        "max_tokens": 1800[cite: 1]
    }[cite: 1]

    try:
        resp = requests.post(DEEPSEEK_CHAT_URL, headers=headers, json=payload, timeout=120)[cite: 1]
        if resp.status_code == 200:[cite: 1]
            res_json = resp.json()[cite: 1]
            usage = res_json.get("usage", {})[cite: 1]
            st.session_state["tokens_used"] += usage.get("total_tokens", len(transcript) // 4)[cite: 1]
            st.session_state["last_api_call"] = datetime.datetime.now()[cite: 1]

            raw_text = res_json["choices"][0]["message"]["content"].strip()[cite: 1]
            clean_text = re.sub(r"^```(?:json)?\s*", "", raw_text)[cite: 1]
            clean_text = re.sub(r"\s*```$", "", clean_text).strip()[cite: 1]
            match = re.search(r"\{.*\}", clean_text, re.DOTALL)[cite: 1]
            data = json.loads(match.group(0)) if match else json.loads(clean_text)[cite: 1]
            return normalize_llm_json_to_df(data)[cite: 1]
        else:
            st.warning(f"DeepSeek Notice ({resp.status_code}): {resp.text}")[cite: 1]
    except Exception as e:[cite: 1]
        st.warning(f"DeepSeek connection error: {e}")[cite: 1]

    return None, ""[cite: 1]

def heuristic_non_ai_extraction(transcript):
    sentences = re.split(r'(?<=[.!?]) +', transcript)[cite: 1]
    action_keywords = ['send', 'prepare', 'submit', 'update', 'review', 'check', 'email', 'kailangan', 'gagawin', 'ipapasa', 'provide', 'target', 'ipresent', 'kukunin'][cite: 1]
    date_keywords = ['tomorrow', 'monday', 'tuesday', 'wednesday', 'thursday', 'friday', 'q1', 'q2', 'q3', 'q4', 'january', 'february', 'march', 'april', 'may', 'june', 'july', 'august', 'september', 'october', 'november', 'december', 'bukas', 'deadline'][cite: 1]
    
    table_items = [][cite: 1]
    other_discussions = [][cite: 1]
    
    for i in range(0, len(sentences), 3):[cite: 1]
        chunk = sentences[i:i+3][cite: 1]
        if not chunk:[cite: 1]
            continue[cite: 1]
        chunk_text = " ".join(chunk)[cite: 1]
        
        has_action = any(kw in chunk_text.lower() for kw in action_keywords)[cite: 1]
        has_date = any(kw in chunk_text.lower() for kw in date_keywords)[cite: 1]
        
        if has_action or has_date:[cite: 1]
            action_text = " ".join([s for s in chunk if any(kw in s.lower() for kw in action_keywords)])[cite: 1]
            table_items.append({[cite: 1]
                "Discussion Points": chunk[0].strip() + "...",[cite: 1]
                "Action Plan": action_text.strip() if action_text else "Review discussion for actions",[cite: 1]
                "Indicative Delivery Date": "Check transcript (Date mentioned)" if has_date else "TBD",[cite: 1]
                "Person-in-charge": "Unassigned"[cite: 1]
            })[cite: 1]
        else:
            other_discussions.append(chunk_text)[cite: 1]
            
    if not table_items:[cite: 1]
        table_items = [{[cite: 1]
            "Discussion Points": "Meeting Overview",[cite: 1]
            "Action Plan": "Please review transcript manually.",[cite: 1]
            "Indicative Delivery Date": "TBD",[cite: 1]
            "Person-in-charge": "Unassigned"[cite: 1]
        }][cite: 1]
        
    df = pd.DataFrame(table_items[:10])[cite: 1]
    for col in ["Discussion Points", "Action Plan", "Indicative Delivery Date", "Person-in-charge"]:[cite: 1]
        if col not in df.columns:[cite: 1]
            df[col] = ""[cite: 1]
            
    df = df[["Discussion Points", "Action Plan", "Indicative Delivery Date", "Person-in-charge"]][cite: 1]
    other_text = "\n\n".join(other_discussions[:4])[cite: 1]
    return df, other_text[cite: 1]

def extract_structured_insights(transcript, engine="AI - DeepSeek"):
    progress_bar = st.progress(0, text="Initializing MOM extraction (0%)...")[cite: 1]
    time.sleep(0.2)[cite: 1]
    progress_bar.progress(40, text=f"Translating Taglish conversation & extracting with {engine} (40%)...")[cite: 1]

    if engine == "Non-AI - Python Heuristic":[cite: 1]
        time.sleep(0.5)[cite: 1]
        res_df, res_other = heuristic_non_ai_extraction(transcript)[cite: 1]
        progress_bar.progress(100, text="Extraction completed (100%)!")[cite: 1]
        time.sleep(0.2)[cite: 1]
        progress_bar.empty()[cite: 1]
        return res_df, res_other[cite: 1]

    df, other = extract_with_deepseek(transcript)[cite: 1]
    
    if df is not None and not df.empty:[cite: 1]
        progress_bar.progress(100, text="Finalizing Minutes of the Meeting (100%)...")[cite: 1]
        time.sleep(0.3)[cite: 1]
        progress_bar.empty()[cite: 1]
        return df, other[cite: 1]

    df_fb, other_fb = heuristic_non_ai_extraction(transcript)[cite: 1]
    progress_bar.empty()[cite: 1]
    st.markdown(f"{SVG_ALERT} AI completion request could not be completed. The table below was populated using offline Keyword Heuristics.", unsafe_allow_html=True)[cite: 1]
    return df_fb, other_fb[cite: 1]

def ask_deepseek_question(transcript, question, chat_history):
    if not DEEPSEEK_API_KEY:[cite: 1]
        return "DeepSeek API key is missing. Please check your configuration."[cite: 1]

    headers = {[cite: 1]
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}",[cite: 1]
        "Content-Type": "application/json"[cite: 1]
    }[cite: 1]

    system_prompt = ([cite: 1]
        "You are Ask Echo, an authentic, executive AI assistant for PRIME Philippines. "
        "Answer questions based accurately and concisely on the provided meeting transcript. "
        "Use subtle, clean Markdown with bullet points where appropriate. "
        "If a specific detail is not in the transcript, concisely state that it was not mentioned."
    )

    messages = [{"role": "system", "content": system_prompt}][cite: 1]
    for msg in chat_history[-6:]:[cite: 1]
        messages.append({"role": msg["role"], "content": msg["content"]})[cite: 1]
    messages.append({"role": "user", "content": f"Transcript:\n{transcript[:22000]}\n\nQuestion: {question}"})[cite: 1]

    payload = {[cite: 1]
        "model": "deepseek-chat",[cite: 1]
        "messages": messages,[cite: 1]
        "temperature": 0.2,[cite: 1]
        "max_tokens": 600
    }

    try:
        resp = requests.post(DEEPSEEK_CHAT_URL, headers=headers, json=payload, timeout=60)[cite: 1]
        if resp.status_code == 200:[cite: 1]
            res_json = resp.json()[cite: 1]
            usage = res_json.get("usage", {})[cite: 1]
            st.session_state["tokens_used"] += usage.get("total_tokens", 0)[cite: 1]
            st.session_state["last_api_call"] = datetime.datetime.now()[cite: 1]
            return res_json["choices"][0]["message"]["content"].strip()[cite: 1]
        else:
            return f"Service notice ({resp.status_code}): {resp.text}"[cite: 1]
    except Exception as e:[cite: 1]
        return f"Connection error: {e}"[cite: 1]

def set_cell_shading(cell, color_hex):
    shd = parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{color_hex}"/>')[cite: 1]
    cell._tc.get_or_add_tcPr().append(shd)[cite: 1]

def export_to_word(df, meeting_details, other_discussions):
    template_files = ["MOM_Template.docx", "MOM Template.docx"][cite: 1]
    template_path = next((f for f in template_files if os.path.exists(f)), None)[cite: 1]

    if template_path:[cite: 1]
        doc = Document(template_path)[cite: 1]
    else:
        doc = Document()[cite: 1]
        if os.path.exists("header.png"):[cite: 1]
            for section in doc.sections:[cite: 1]
                section.top_margin = Inches(0.4)[cite: 1]
                section.bottom_margin = Inches(0.4)[cite: 1]
                section.left_margin = Inches(0.75)[cite: 1]
                section.right_margin = Inches(0.75)[cite: 1]
                hp = section.header.paragraphs[0][cite: 1]
                hp.alignment = WD_ALIGN_PARAGRAPH.CENTER[cite: 1]
                hp.add_run().add_picture("header.png", width=Inches(7.0))[cite: 1]
                if os.path.exists("footer.png"):[cite: 1]
                    fp = section.footer.paragraphs[0][cite: 1]
                    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER[cite: 1]
                    fp.add_run().add_picture("footer.png", width=Inches(7.0))[cite: 1]

    for section in doc.sections:[cite: 1]
        section.top_margin = Inches(0.4)[cite: 1]
        section.bottom_margin = Inches(0.4)[cite: 1]
        section.left_margin = Inches(0.75)[cite: 1]
        section.right_margin = Inches(0.75)[cite: 1]

    p_title = doc.add_paragraph()[cite: 1]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER[cite: 1]
    p_title.paragraph_format.space_before = Pt(0)[cite: 1]
    p_title.paragraph_format.space_after = Pt(2)[cite: 1]
    r_title = p_title.add_run("MINUTES OF THE MEETING")[cite: 1]
    r_title.bold = True[cite: 1]
    r_title.underline = True[cite: 1]
    r_title.font.name = "Arial"[cite: 1]
    r_title.font.size = Pt(11)[cite: 1]

    company_target = meeting_details.get("external_attendees", [])[cite: 1]
    primary_client_rep = company_target[0] if company_target else meeting_details.get("company_name", "").strip() or "CLIENT"[cite: 1]
    p_sub = doc.add_paragraph()[cite: 1]
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER[cite: 1]
    p_sub.paragraph_format.space_after = Pt(12)[cite: 1]
    r_sub = p_sub.add_run(f"PRIME PHILIPPINES & {primary_client_rep.upper()}")[cite: 1]
    r_sub.bold = True[cite: 1]
    r_sub.font.name = "Arial"[cite: 1]
    r_sub.font.size = Pt(11)[cite: 1]

    date_str = meeting_details.get("date", "____________")[cite: 1]
    time_str = meeting_details.get("time_range", "")[cite: 1]
    full_date = f"Date: {date_str}" + (f", {time_str}" if time_str.strip() else "")[cite: 1]
    
    p_date = doc.add_paragraph(full_date)[cite: 1]
    p_date.paragraph_format.space_after = Pt(2)[cite: 1]
    for r in p_date.runs:[cite: 1]
        r.font.name = "Arial"[cite: 1]
        r.font.size = Pt(10)[cite: 1]

    p_loc = doc.add_paragraph(f"Location: {meeting_details.get('location', '____________')}")[cite: 1]
    p_loc.paragraph_format.space_after = Pt(2)[cite: 1]
    for r in p_loc.runs:[cite: 1]
        r.font.name = "Arial"[cite: 1]
        r.font.size = Pt(10)[cite: 1]

    prime_atts = meeting_details.get("prime_attendees", [])[cite: 1]
    ext_atts = meeting_details.get("external_attendees", [])[cite: 1]
    
    p_att = doc.add_paragraph()[cite: 1]
    p_att.paragraph_format.space_after = Pt(2)[cite: 1]
    p_att.paragraph_format.tab_stops.add_tab_stop(Inches(1.35), WD_TAB_ALIGNMENT.LEFT)[cite: 1]
    r_att_label = p_att.add_run("Attended by:")[cite: 1]
    r_att_label.font.name = "Arial"[cite: 1]
    r_att_label.font.size = Pt(10)[cite: 1]
    
    first_attendee = True[cite: 1]
    if ext_atts:[cite: 1]
        for att in ext_atts:[cite: 1]
            if not att.strip():[cite: 1]
                continue[cite: 1]
            p = p_att if first_attendee else doc.add_paragraph()[cite: 1]
            p.paragraph_format.space_after = Pt(2)[cite: 1]
            if not first_attendee:[cite: 1]
                p.paragraph_format.left_indent = Inches(1.35)[cite: 1]
            else:
                p.add_run("\t")[cite: 1]
            comp_label = f", {meeting_details.get('company_name')}" if meeting_details.get('company_name') else ""[cite: 1]
            r = p.add_run(f"{att}{comp_label}")[cite: 1]
            r.font.name = "Arial"[cite: 1]
            r.font.size = Pt(10)[cite: 1]
            first_attendee = False[cite: 1]

    if prime_atts:[cite: 1]
        for att in prime_atts:[cite: 1]
            p = p_att if first_attendee else doc.add_paragraph()[cite: 1]
            p.paragraph_format.space_after = Pt(2)[cite: 1]
            if not first_attendee:[cite: 1]
                p.paragraph_format.left_indent = Inches(1.35)[cite: 1]
            else:
                p.add_run("\t")[cite: 1]
            r = p.add_run(f"{att} – PRIME Philippines")[cite: 1]
            r.font.name = "Arial"[cite: 1]
            r.font.size = Pt(10)[cite: 1]
            first_attendee = False[cite: 1]

    p_line = doc.add_paragraph()[cite: 1]
    p_line.paragraph_format.space_before = Pt(4)[cite: 1]
    p_line.paragraph_format.space_after = Pt(6)[cite: 1]
    r_line = p_line.add_run("_________________________________________________________________________________")[cite: 1]
    r_line.font.name = "Arial"[cite: 1]
    r_line.font.color.rgb = RGBColor(160, 160, 160)[cite: 1]

    client_display = meeting_details.get('company_name', '').strip() or "the Client"[cite: 1]
    p_intro = doc.add_paragraph([cite: 1]
        f"During the meeting held last {date_str}, PRIME Philippines, represented by the attendee/s shown above, "[cite: 1]
        f"met with {client_display} to discuss opportunities for collaboration."[cite: 1]
    )[cite: 1]
    p_intro.paragraph_format.space_after = Pt(10)[cite: 1]
    for r in p_intro.runs:[cite: 1]
        r.font.name = "Arial"[cite: 1]
        r.font.size = Pt(9.5)[cite: 1]

    table = doc.add_table(rows=len(df)+1, cols=4)[cite: 1]
    table.alignment = WD_TABLE_ALIGNMENT.CENTER[cite: 1]
    table.style = "Table Grid"[cite: 1]
    table.autofit = False[cite: 1]
    table.allow_autofit = False[cite: 1]

    col_widths = [Inches(2.5), Inches(2.2), Inches(1.1), Inches(1.2)][cite: 1]

    headers = ["Discussion Points", "Action Plan", "Indicative Delivery Date", "Person-in-charge"][cite: 1]
    for i, header in enumerate(headers):[cite: 1]
        cell = table.rows[0].cells[i][cite: 1]
        cell.width = col_widths[i][cite: 1]
        cell.text = header[cite: 1]
        set_cell_shading(cell, "FFFF00")[cite: 1]
        p = cell.paragraphs[0][cite: 1]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER[cite: 1]
        if p.runs:[cite: 1]
            p.runs[0].font.bold = True[cite: 1]
            p.runs[0].font.size = Pt(9)[cite: 1]
            p.runs[0].font.name = "Arial"[cite: 1]

    for i, row in df.iterrows():[cite: 1]
        cells = table.rows[i+1].cells[cite: 1]
        cells[0].text = f"{i+1}. {str(row.get('Discussion Points', ''))}"[cite: 1]
        cells[1].text = str(row.get("Action Plan", ""))[cite: 1]
        cells[2].text = str(row.get("Indicative Delivery Date", ""))[cite: 1]
        cells[3].text = str(row.get("Person-in-charge", ""))[cite: 1]
        for c_idx, cell in enumerate(cells):[cite: 1]
            cell.width = col_widths[c_idx][cite: 1]
            p = cell.paragraphs[0][cite: 1]
            if c_idx in [2, 3]:[cite: 1]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER[cite: 1]
            if p.runs:[cite: 1]
                p.runs[0].font.size = Pt(8.5)[cite: 1]
                p.runs[0].font.name = "Arial"[cite: 1]

    doc.add_paragraph()[cite: 1]
    p_note = doc.add_paragraph("*Note: The indicative delivery date serves as reference point and still subject to changes. Furthermore, it depends on the progress of both parties.")[cite: 1]
    p_note.paragraph_format.space_after = Pt(8)[cite: 1]
    p_note.runs[0].font.italic = True[cite: 1]
    p_note.runs[0].font.name = "Arial"[cite: 1]
    p_note.runs[0].font.size = Pt(8)[cite: 1]

    if other_discussions.strip():[cite: 1]
        p_od_head = doc.add_paragraph()[cite: 1]
        p_od_head.paragraph_format.space_before = Pt(6)[cite: 1]
        p_od_head.paragraph_format.space_after = Pt(4)[cite: 1]
        r_od_head = p_od_head.add_run("Other Discussions:")[cite: 1]
        r_od_head.bold = True[cite: 1]
        r_od_head.font.size = Pt(10)[cite: 1]
        r_od_head.font.name = "Arial"[cite: 1]
        
        p_od = doc.add_paragraph(other_discussions)[cite: 1]
        p_od.paragraph_format.space_after = Pt(12)[cite: 1]
        for r in p_od.runs:[cite: 1]
            r.font.name = "Arial"[cite: 1]
            r.font.size = Pt(9.5)[cite: 1]

    p_prep_label = doc.add_paragraph("Prepared by:")[cite: 1]
    p_prep_label.paragraph_format.space_before = Pt(12)[cite: 1]
    p_prep_label.paragraph_format.space_after = Pt(2)[cite: 1]
    p_prep_label.runs[0].font.name = "Arial"[cite: 1]
    p_prep_label.runs[0].font.bold = True[cite: 1]
    p_prep_label.runs[0].font.size = Pt(9.5)[cite: 1]

    p_prep_line = doc.add_paragraph("_______________________________")[cite: 1]
    p_prep_line.paragraph_format.space_after = Pt(2)[cite: 1]
    p_prep_line.runs[0].font.name = "Arial"[cite: 1]

    prep_name = meeting_details.get("prep_name", "").strip() or "____________________"[cite: 1]
    prep_desig = meeting_details.get("prep_desig", "").strip() or "PRIME Philippines"[cite: 1]
    p_prep_info = doc.add_paragraph(f"{prep_name}\n{prep_desig}")[cite: 1]
    p_prep_info.paragraph_format.space_after = Pt(12)[cite: 1]
    for r in p_prep_info.runs:[cite: 1]
        r.font.name = "Arial"[cite: 1]
        r.font.size = Pt(9.5)[cite: 1]

    p_conf_label = doc.add_paragraph("Confirmed by:")[cite: 1]
    p_conf_label.paragraph_format.space_after = Pt(2)[cite: 1]
    p_conf_label.runs[0].font.name = "Arial"[cite: 1]
    p_conf_label.runs[0].font.bold = True[cite: 1]
    p_conf_label.runs[0].font.size = Pt(9.5)[cite: 1]

    p_conf_line = doc.add_paragraph("_______________________________")[cite: 1]
    p_conf_line.paragraph_format.space_after = Pt(2)[cite: 1]
    p_conf_line.runs[0].font.name = "Arial"[cite: 1]

    conf_name = meeting_details.get("conf_name", "").strip() or (ext_atts[0] if ext_atts else "____________________")[cite: 1]
    conf_desig = meeting_details.get("conf_desig", "").strip() or (meeting_details.get("company_name", "").strip() or "Client")[cite: 1]
    p_conf_info = doc.add_paragraph(f"{conf_name}\n{conf_desig}")[cite: 1]
    p_conf_info.paragraph_format.space_after = Pt(6)[cite: 1]
    for r in p_conf_info.runs:[cite: 1]
        r.font.name = "Arial"[cite: 1]
        r.font.size = Pt(9.5)[cite: 1]

    bio = BytesIO()[cite: 1]
    doc.save(bio)[cite: 1]
    bio.seek(0)[cite: 1]
    return bio[cite: 1]

def export_to_pdf(df, meeting_details, other_discussions):
    buffer = BytesIO()[cite: 1]
    doc = SimpleDocTemplate([cite: 1]
        buffer,[cite: 1]
        pagesize=letter,[cite: 1]
        leftMargin=0.6 * inch,[cite: 1]
        rightMargin=0.6 * inch,[cite: 1]
        topMargin=0.5 * inch,[cite: 1]
        bottomMargin=0.5 * inch[cite: 1]
    )[cite: 1]
    story = [][cite: 1]
    styles = getSampleStyleSheet()[cite: 1]

    style_title = ParagraphStyle([cite: 1]
        'DocTitle',[cite: 1]
        parent=styles['Normal'],[cite: 1]
        fontName='Helvetica-Bold',[cite: 1]
        fontSize=11.5,[cite: 1]
        alignment=1,[cite: 1]
        spaceAfter=2[cite: 1]
    )[cite: 1]
    company_target = meeting_details.get("external_attendees", [])[cite: 1]
    primary_client_rep = company_target[0] if company_target else meeting_details.get("company_name", "").strip() or "CLIENT"[cite: 1]
    style_subtitle = ParagraphStyle([cite: 1]
        'DocSubTitle',[cite: 1]
        parent=styles['Normal'],[cite: 1]
        fontName='Helvetica-Bold',[cite: 1]
        fontSize=10.5,[cite: 1]
        alignment=1,[cite: 1]
        spaceAfter=10[cite: 1]
    )[cite: 1]
    style_body = ParagraphStyle([cite: 1]
        'DocBody',[cite: 1]
        parent=styles['Normal'],[cite: 1]
        fontName='Helvetica',[cite: 1]
        fontSize=9,[cite: 1]
        leading=12,[cite: 1]
        spaceAfter=3[cite: 1]
    )[cite: 1]
    style_th = ParagraphStyle([cite: 1]
        'TableHead',[cite: 1]
        parent=styles['Normal'],[cite: 1]
        fontName='Helvetica-Bold',[cite: 1]
        fontSize=8.5,[cite: 1]
        leading=10,[cite: 1]
        alignment=1[cite: 1]
    )[cite: 1]
    style_td = ParagraphStyle([cite: 1]
        'TableData',[cite: 1]
        parent=styles['Normal'],[cite: 1]
        fontName='Helvetica',[cite: 1]
        fontSize=8,[cite: 1]
        leading=10[cite: 1]
    )[cite: 1]
    style_td_center = ParagraphStyle([cite: 1]
        'TableDataCenter',[cite: 1]
        parent=styles['Normal'],[cite: 1]
        fontName='Helvetica',[cite: 1]
        fontSize=8,[cite: 1]
        leading=10,[cite: 1]
        alignment=1[cite: 1]
    )[cite: 1]

    if os.path.exists("header.png"):[cite: 1]
        story.append(Image("header.png", width=6.8 * inch, height=0.75 * inch))[cite: 1]
        story.append(Spacer(1, 6))[cite: 1]

    story.append(Paragraph("<u>MINUTES OF THE MEETING</u>", style_title))[cite: 1]
    story.append(Paragraph(f"PRIME PHILIPPINES & {primary_client_rep.upper()}", style_subtitle))[cite: 1]

    date_str = meeting_details.get("date", "____________")[cite: 1]
    time_str = meeting_details.get("time_range", "")[cite: 1]
    full_date = f"<b>Date:</b> {date_str}" + (f", {time_str}" if time_str.strip() else "")[cite: 1]
    story.append(Paragraph(full_date, style_body))[cite: 1]
    story.append(Paragraph(f"<b>Location:</b> {meeting_details.get('location', '____________')}", style_body))[cite: 1]

    prime_atts = meeting_details.get("prime_attendees", [])[cite: 1]
    ext_atts = meeting_details.get("external_attendees", [])[cite: 1]
    att_list = [][cite: 1]
    if ext_atts:[cite: 1]
        for att in ext_atts:[cite: 1]
            if att.strip():[cite: 1]
                comp_label = f", {meeting_details.get('company_name')}" if meeting_details.get('company_name') else ""[cite: 1]
                att_list.append(f"{att}{comp_label}")[cite: 1]
    if prime_atts:[cite: 1]
        for att in prime_atts:[cite: 1]
            att_list.append(f"{att} – PRIME Philippines")[cite: 1]

    if att_list:[cite: 1]
        story.append(Paragraph(f"<b>Attended by:</b>&nbsp;&nbsp;&nbsp;&nbsp;{att_list[0]}", style_body))[cite: 1]
        for a in att_list[1:]:[cite: 1]
            story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;{a}", style_body))[cite: 1]

    story.append(Spacer(1, 4))[cite: 1]
    client_display = meeting_details.get('company_name', '').strip() or "the Client"[cite: 1]
    story.append(Paragraph([cite: 1]
        f"During the meeting held last {date_str}, PRIME Philippines, represented by the attendee/s shown above, "[cite: 1]
        f"met with {client_display} to discuss opportunities for collaboration.",[cite: 1]
        style_body[cite: 1]
    ))[cite: 1]
    story.append(Spacer(1, 6))[cite: 1]

    table_data = [[[cite: 1]
        Paragraph("<b>Discussion Points</b>", style_th),[cite: 1]
        Paragraph("<b>Action Plan</b>", style_th),[cite: 1]
        Paragraph("<b>Indicative Delivery Date</b>", style_th),[cite: 1]
        Paragraph("<b>Person-in-charge</b>", style_th)[cite: 1]
    ]][cite: 1]

    for i, row in df.iterrows():[cite: 1]
        table_data.append([[cite: 1]
            Paragraph(f"{i+1}. {str(row.get('Discussion Points', ''))}", style_td),[cite: 1]
            Paragraph(str(row.get("Action Plan", "")), style_td),[cite: 1]
            Paragraph(str(row.get("Indicative Delivery Date", "")), style_td_center),[cite: 1]
            Paragraph(str(row.get("Person-in-charge", "")), style_td_center)[cite: 1]
        ])[cite: 1]

    col_widths = [2.4 * inch, 2.3 * inch, 1.1 * inch, 1.0 * inch][cite: 1]
    t = Table(table_data, colWidths=col_widths, repeatRows=1)[cite: 1]
    t.setStyle(TableStyle([[cite: 1]
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#FFFF00')),[cite: 1]
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),[cite: 1]
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),[cite: 1]
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),[cite: 1]
        ('TOPPADDING', (0, 0), (-1, -1), 4),[cite: 1]
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),[cite: 1]
    ]))[cite: 1]
    story.append(t)[cite: 1]

    note_style = ParagraphStyle('Note', parent=styles['Normal'], fontName='Helvetica-Oblique', fontSize=7.5, leading=9, spaceBefore=4)[cite: 1]
    story.append(Paragraph("*Note: The indicative delivery date serves as reference point and still subject to changes. Furthermore, it depends on the progress of both parties.", note_style))[cite: 1]

    if other_discussions.strip():[cite: 1]
        story.append(Spacer(1, 6))[cite: 1]
        story.append(Paragraph("<b>Other Discussions:</b>", style_body))[cite: 1]
        story.append(Paragraph(other_discussions, style_body))[cite: 1]

    story.append(Spacer(1, 8))[cite: 1]
    prep_name = meeting_details.get("prep_name", "").strip() or "____________________"[cite: 1]
    prep_desig = meeting_details.get("prep_desig", "").strip() or "PRIME Philippines"[cite: 1]
    conf_name = meeting_details.get("conf_name", "").strip() or (ext_atts[0] if ext_atts else "____________________")[cite: 1]
    conf_desig = meeting_details.get("conf_desig", "").strip() or (meeting_details.get("company_name", "").strip() or "Client")[cite: 1]

    sign_data = [[cite: 1]
        [Paragraph("<b>Prepared by:</b>", style_body), Paragraph("<b>Confirmed by:</b>", style_body)],[cite: 1]
        [Paragraph("_______________________________", style_body), Paragraph("_______________________________", style_body)],[cite: 1]
        [Paragraph(f"{prep_name}<br/>{prep_desig}", style_body), Paragraph(f"{conf_name}<br/>{conf_desig}", style_body)][cite: 1]
    ][cite: 1]
    sign_table = Table(sign_data, colWidths=[3.4 * inch, 3.4 * inch])[cite: 1]
    sign_table.setStyle(TableStyle([[cite: 1]
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),[cite: 1]
        ('LEFTPADDING', (0, 0), (-1, -1), 0),[cite: 1]
        ('RIGHTPADDING', (0, 0), (-1, -1), 0),[cite: 1]
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),[cite: 1]
    ]))[cite: 1]
    story.append(sign_table)[cite: 1]

    if os.path.exists("footer.png"):[cite: 1]
        story.append(Spacer(1, 8))[cite: 1]
        story.append(Image("footer.png", width=6.8 * inch, height=0.65 * inch))[cite: 1]

    doc.build(story)[cite: 1]
    buffer.seek(0)[cite: 1]
    return buffer[cite: 1]

# ========== STREAMLIT UI SETUP ==========
st.markdown(CUSTOM_CSS, unsafe_allow_html=True)[cite: 1]

# Top Bar Fixed Header
st.markdown("""
<div class="echo-topbar-wrapper">
 <h1 class="echo-title">Project <span>Echo</span></h1>
</div>
""", unsafe_allow_html=True)[cite: 1]

# ---- TOP ROW: Upload Section (Left) and Meeting Details (Right) in 2 Separate Containers ----
col_upload, col_details = st.columns(2)

# LEFT CONTAINER: Audio & Text Upload Section
with col_upload:
    with st.container(border=True):
        st.markdown('<h3>Input & Transcription</h3>', unsafe_allow_html=True)
        
        tab_upload, tab_record, tab_text = st.tabs(["Upload Audio", "Record Audio", "Upload Text"])

        # TAB 1: UPLOAD AUDIO
        with tab_upload:
            uploaded_file = st.file_uploader(
                "Upload audio file (200MB limit supported)",
                type=["wav", "mp3", "m4a", "ogg", "flac", "mp4", "webm"],
                help="Audio uploads up to 200MB are supported."
            )
            if uploaded_file:
                st.write("")
                if st.button("Transcribe Audio", key="btn_tx_upload"):
                    p_bar = st.progress(0, text="Initializing audio pipeline (0%)...")
                    p_status = st.empty()
                    transcript = transcribe_audio_pipeline(uploaded_file.read(), uploaded_file.name, p_bar, p_status)
                    p_bar.empty()
                    p_status.empty()
                    if transcript:
                        st.session_state["transcript"] = transcript
                        st.session_state["df"] = pd.DataFrame(columns=["Discussion Points", "Action Plan", "Indicative Delivery Date", "Person-in-charge"])
                        st.session_state["other_discussions"] = ""
                        st.session_state["chat_history"] = []
                        st.rerun()

        # TAB 2: RECORD AUDIO
        with tab_record:
            recorded_audio = st.audio_input("Record audio directly", label_visibility="collapsed")
            if recorded_audio:
                rec_bytes = recorded_audio.read()
                r_btn1, r_btn2 = st.columns(2)
                with r_btn1:
                    st.download_button(label="Save Recording (.wav)", data=rec_bytes, file_name=f"Recording_{datetime.date.today().strftime('%Y%m%d')}.wav", mime="audio/wav", use_container_width=True)
                with r_btn2:
                    if st.button("Transcribe Audio", key="btn_tx_record"):
                        p_bar = st.progress(0, text="Initializing audio pipeline (0%)...")
                        p_status = st.empty()
                        transcript = transcribe_audio_pipeline(rec_bytes, "recording.wav", p_bar, p_status)
                        p_bar.empty()
                        p_status.empty()
                        if transcript:
                            st.session_state["transcript"] = transcript
                            st.session_state["df"] = pd.DataFrame(columns=["Discussion Points", "Action Plan", "Indicative Delivery Date", "Person-in-charge"])
                            st.session_state["other_discussions"] = ""
                            st.session_state["chat_history"] = []
                            st.rerun()

        # TAB 3: TEXT UPLOAD
        with tab_text:
            uploaded_text_file = st.file_uploader("Upload Document (.txt, .docx, .pdf)", type=["txt", "docx", "pdf"])
            pasted_text = st.text_area("Or Paste Transcript Here", height=95, placeholder="Paste transcript text directly here...")
            if st.button("Process Text", key="btn_tx_text"):
                p_bar = st.progress(0, text="Extracting document text (0%)...")
                time.sleep(0.2)
                p_bar.progress(50, text="Reading document stream (50%)...")
                extracted_str = ""
                if uploaded_text_file:
                    extracted_str = extract_text_from_file(uploaded_text_file)
                if pasted_text and pasted_text.strip():
                    extracted_str += "\n" + pasted_text.strip()
                
                p_bar.progress(100, text="Document processed (100%)!")
                time.sleep(0.2)
                p_bar.empty()
                if extracted_str.strip():
                    st.session_state["transcript"] = extracted_str.strip()
                    st.session_state["df"] = pd.DataFrame(columns=["Discussion Points", "Action Plan", "Indicative Delivery Date", "Person-in-charge"])
                    st.session_state["other_discussions"] = ""
                    st.session_state["chat_history"] = []
                    st.rerun()
                else:
                    st.warning("Please upload a file or paste text to proceed.")

# RIGHT CONTAINER: Meeting Details Card
with col_details:
    with st.container(border=True):
        head_col1, head_col2 = st.columns([9.0, 1.0])
        with head_col1:
            st.markdown('<h3>Meeting Details</h3>', unsafe_allow_html=True)
        with head_col2:
            if st.button("", key="card_settings_btn", help="Open MoM Generation Engine & Token Diagnostics"):
                st.session_state["show_settings"] = not st.session_state["show_settings"]
                st.rerun()

        # Settings Drawer
        if st.session_state["show_settings"]:
            with st.expander("Settings & Engine Diagnostics", expanded=True):
                set_col1, set_col2 = st.columns(2)
                with set_col1:
                    engine_options = ["AI - DeepSeek", "Non-AI - Python Heuristic"]
                    selected_eng = st.selectbox(
                        "MoM Generation Engine",
                        options=engine_options,
                        index=engine_options.index(st.session_state["selected_engine"]) if st.session_state["selected_engine"] in engine_options else 0
                    )
                    st.session_state["selected_engine"] = selected_eng

                    if st.button("Regenerate MOM", key="btn_regen_mom"):
                        if st.session_state["transcript"]:
                            extracted_df, other_disc = extract_structured_insights(st.session_state["transcript"], selected_eng)
                            if not extracted_df.empty:
                                st.session_state["df"] = extracted_df
                                st.session_state["other_discussions"] = other_disc
                                st.rerun()
                with set_col2:
                    st.markdown("**Diagnostics**")
                    st.write(f"• **Session Tokens:** `{st.session_state['tokens_used']:,}`")
                    if st.session_state["last_api_call"]:
                        last_call = st.session_state["last_api_call"]
                        st.write(f"• **Last Call:** `{last_call.strftime('%I:%M:%S %p')}`")
                    st.write(f"• **Status:** `Active & Ready`")
            st.markdown("---")

        # Row 1: Date, Location Preset & Input
        r1_c1, r1_c2 = st.columns([1.2, 2.0])
        with r1_c1:
            meeting_date = st.date_input("Date", value=datetime.date(2026, 8, 25))
        with r1_c2:
            loc_preset = st.selectbox("Location Preset", options=LOCATION_PRESETS, index=0)
            custom_loc = st.text_input("Location", value="", placeholder="e.g. Boardroom", label_visibility="collapsed")
            meeting_location = custom_loc.strip() if custom_loc.strip() else ("" if loc_preset == LOCATION_PRESETS[0] else loc_preset)

        # Row 2: Time Pickers
        r2_c1, r2_c2 = st.columns(2)
        with r2_c1:
            st.markdown("<p style='font-size:0.85rem; margin-bottom:0.2rem; color:#333; font-weight:500;'>Start Time</p>", unsafe_allow_html=True)
            sc1, sc2, sc3 = st.columns([1, 1, 1.2])
            sh = sc1.selectbox("SH", [f"{i:02d}" for i in range(1,13)], key="sh", label_visibility="collapsed")
            sm = sc2.selectbox("SM", [f"{i:02d}" for i in range(0,60,5)], key="sm", label_visibility="collapsed")
            sap = sc3.selectbox("SAP", ["AM", "PM"], key="sap", label_visibility="collapsed")
            start_str = f"{sh}:{sm} {sap}"
        with r2_c2:
            st.markdown("<p style='font-size:0.85rem; margin-bottom:0.2rem; color:#333; font-weight:500;'>End Time</p>", unsafe_allow_html=True)
            ec1, ec2, ec3 = st.columns([1, 1, 1.2])
            eh = ec1.selectbox("EH", [f"{i:02d}" for i in range(1,13)], key="eh", label_visibility="collapsed")
            em = ec2.selectbox("EM", [f"{i:02d}" for i in range(0,60,5)], key="em", label_visibility="collapsed")
            eap = ec3.selectbox("EAP", ["AM", "PM"], key="eap", label_visibility="collapsed")
            end_str = f"{eh}:{em} {eap}"

        # Row 3: Attendees & Parties
        r3_c1, r3_c2 = st.columns(2)
        with r3_c1:
            client_name = st.text_input("Client / Company", value="", placeholder="XYZ Company")
            selected_crd = st.multiselect("CRD Team Attendees", options=CRD_MEMBERS, default=[])
        with r3_c2:
            ext_attendees_raw = st.text_input("External Attendees", value="", placeholder="e.g. Mr. ABCD, Jane Doe")
            prep_col, conf_col = st.columns(2)
            with prep_col:
                prep_name = st.text_input("Prepared By", value="", placeholder="Name")
                prep_desig = st.text_input("Prep Designation", value="", placeholder="Designation")
            with conf_col:
                conf_name = st.text_input("Confirmed By", value="", placeholder="Name")
                conf_desig = st.text_input("Conf Designation", value="", placeholder="Designation")

# ---- Step 2: Full Transcript & Ask Echo Side-by-Side ----
if st.session_state["transcript"]:
    row_left, row_right = st.columns(2)
    
    # LEFT CONTAINER: Full Transcript
    with row_left:
        with st.container(border=True):
            ft_head, ft_btn1, ft_btn2 = st.columns([6, 2, 2])
            with ft_head:
                st.markdown('<h3 style="margin-top:0.2rem;">Full Transcript</h3>', unsafe_allow_html=True)[cite: 1]
            with ft_btn1:
                copy_html = f"""
                <!DOCTYPE html>
                <html>
                <head>
                <style>
                body {{ margin: 0; padding: 0; font-family: 'Montserrat', sans-serif; }}
                button {{
                    width: 100%;
                    height: 34px;
                    background-color: #222222;
                    color: #FFFFFF;
                    border: 1px solid #444444;
                    border-radius: 50px;
                    font-size: 0.82rem;
                    font-weight: 500;
                    cursor: pointer;
                    transition: all 0.25s ease;
                    display: flex;
                    align-items: center;
                    justify-content: center;
                }}
                button:hover {{
                    border-color: #D4AF37;
                    color: #D4AF37;
                    background-color: #1A1A1A;
                }}
                </style>
                </head>
                <body>
                    <button id="copy-btn">Copy Text</button>
                    <script>
                    document.getElementById("copy-btn").addEventListener("click", function() {{
                        navigator.clipboard.writeText({json.dumps(st.session_state["transcript"])}).then(function() {{
                            document.getElementById("copy-btn").innerText = "Copied";
                            setTimeout(() => document.getElementById("copy-btn").innerText = "Copy Text", 2000);
                        }});
                    }});
                    </script>
                </body>
                </html>
                """[cite: 1]
                components.html(copy_html, height=34)[cite: 1]
            with ft_btn2:
                st.download_button([cite: 1]
                    label="Download",[cite: 1]
                    data=st.session_state["transcript"],[cite: 1]
                    file_name=f"Transcript_{meeting_date.strftime('%Y%m%d')}.txt",[cite: 1]
                    mime="text/plain",[cite: 1]
                    use_container_width=True[cite: 1]
                )[cite: 1]
            
            st.text_area([cite: 1]
                "Transcript Content", [cite: 1]
                st.session_state["transcript"], [cite: 1]
                height=380,  [cite: 1]
                label_visibility="collapsed"[cite: 1]
            )[cite: 1]
            
            if st.session_state["df"].empty:[cite: 1]
                st.write("")[cite: 1]
                if st.button("Generate MOM", key="btn_gen_mom"):[cite: 1]
                    extracted_df, other_disc = extract_structured_insights(st.session_state["transcript"], st.session_state["selected_engine"])[cite: 1]
                    if not extracted_df.empty:[cite: 1]
                        st.session_state["df"] = extracted_df[cite: 1]
                        st.session_state["other_discussions"] = other_disc[cite: 1]
                        st.rerun()[cite: 1]

    # RIGHT CONTAINER: Ask Echo (AI on Left, User on Right, Sophisticated Theme)
    with row_right:
        with st.container(border=True):
            st.markdown('<h3 style="margin-top:0.2rem;">Ask Echo</h3>', unsafe_allow_html=True)[cite: 1]
            st.caption("Ask specific questions regarding action items, timelines, deliverables, or remarks.")[cite: 1]
            
            # Chat history container with Minimalist/Mature Styling
            st.markdown('<div class="chat-container">', unsafe_allow_html=True)[cite: 1]
            if not st.session_state["chat_history"]:[cite: 1]
                st.markdown([cite: 1]
                    '<div class="chat-bubble-ai-wrap">'
                    '<div class="chat-bubble-ai"><div class="ai-header">Echo</div>How may I assist you regarding this meeting transcript?</div>'
                    '</div>',
                    unsafe_allow_html=True
                )
            else:
                for msg in st.session_state["chat_history"]:[cite: 1]
                    if msg["role"] == "assistant":[cite: 1]
                        # AI Bubble (Left-Aligned, Elegant Off-White & Gold)
                        formatted_content = msg["content"].replace("\n", "<br>")
                        st.markdown(
                            f'<div class="chat-bubble-ai-wrap">'
                            f'<div class="chat-bubble-ai"><div class="ai-header">Echo</div>{formatted_content}</div>'
                            f'</div>',
                            unsafe_allow_html=True
                        )
                    else:
                        # User Bubble (Right-Aligned, Warm Muted Pill)
                        st.markdown(
                            f'<div class="chat-bubble-user-wrap">'
                            f'<div class="chat-bubble-user">{msg["content"]}</div>'
                            f'</div>',
                            unsafe_allow_html=True
                        )[cite: 1]
            st.markdown('</div>', unsafe_allow_html=True)[cite: 1]
            
            # Chat input
            if prompt := st.chat_input("Ask Echo a question..."):[cite: 1]
                st.session_state["chat_history"].append({"role": "user", "content": prompt})[cite: 1]
                with st.spinner("Analyzing transcript..."):[cite: 1]
                    answer = ask_deepseek_question(st.session_state["transcript"], prompt, st.session_state["chat_history"])[cite: 1]
                st.session_state["chat_history"].append({"role": "assistant", "content": answer})[cite: 1]
                st.rerun()[cite: 1]

# ---- Step 3: Minutes of Meeting Editor ----
if not st.session_state["df"].empty:[cite: 1]
    with st.container(border=True):[cite: 1]
        st.markdown('<h3>Minutes of Meeting Editor</h3>', unsafe_allow_html=True)[cite: 1]
        
        st.markdown([cite: 1]
            "<p style='font-size:0.85rem; color:#666; margin-bottom: 0.75rem;'>"[cite: 1]
            "<i>*Note: Each discussion item is rendered as a clean card with auto-wrapping text boxes. Edit fields inline directly.</i></p>", [cite: 1]
            unsafe_allow_html=True[cite: 1]
        )[cite: 1]
        
        df = st.session_state["df"].copy().reset_index(drop=True)[cite: 1]
        
        row_to_delete = None[cite: 1]
        for idx in range(len(df)):[cite: 1]
            with st.container(border=True):[cite: 1]
                c_disc, c_act, c_date, c_pic, c_del = st.columns([3.2, 3.2, 1.8, 1.8, 0.6])[cite: 1]
                
                with c_disc:[cite: 1]
                    st.markdown('<span class="playfair-label">Discussion Points</span>', unsafe_allow_html=True)[cite: 1]
                    st.text_area([cite: 1]
                        "DP",[cite: 1]
                        value=str(df.at[idx, "Discussion Points"]),[cite: 1]
                        key=f"dp_{idx}",[cite: 1]
                        height=75,[cite: 1]
                        label_visibility="collapsed"[cite: 1]
                    )[cite: 1]
                with c_act:[cite: 1]
                    st.markdown('<span class="playfair-label">Action Plan</span>', unsafe_allow_html=True)[cite: 1]
                    st.text_area([cite: 1]
                        "AP",[cite: 1]
                        value=str(df.at[idx, "Action Plan"]),[cite: 1]
                        key=f"ap_{idx}",[cite: 1]
                        height=75,[cite: 1]
                        label_visibility="collapsed"[cite: 1]
                    )[cite: 1]
                with c_date:[cite: 1]
                    st.markdown('<span class="playfair-label">Delivery Date</span>', unsafe_allow_html=True)[cite: 1]
                    st.text_area([cite: 1]
                        "DD",[cite: 1]
                        value=str(df.at[idx, "Indicative Delivery Date"]),[cite: 1]
                        key=f"date_{idx}",[cite: 1]
                        height=75,[cite: 1]
                        label_visibility="collapsed"[cite: 1]
                    )[cite: 1]
                with c_pic:[cite: 1]
                    st.markdown('<span class="playfair-label">Person-in-charge</span>', unsafe_allow_html=True)[cite: 1]
                    st.text_area([cite: 1]
                        "PIC",[cite: 1]
                        value=str(df.at[idx, "Person-in-charge"]),[cite: 1]
                        key=f"pic_{idx}",[cite: 1]
                        height=75,[cite: 1]
                        label_visibility="collapsed"[cite: 1]
                    )[cite: 1]
                with c_del:[cite: 1]
                    st.write("<div style='height: 38px;'></div>", unsafe_allow_html=True)[cite: 1]
                    if st.button("Delete", key=f"del_{idx}", help=f"Remove item {idx+1}"):[cite: 1]
                        row_to_delete = idx[cite: 1]
        
        # Handle Deletion
        if row_to_delete is not None:[cite: 1]
            df = df.drop(index=row_to_delete).reset_index(drop=True)[cite: 1]
            st.session_state["df"] = df[cite: 1]
            st.rerun()[cite: 1]
        
        # Collect updated values using robust list aggregation
        rows_data = [][cite: 1]
        for idx in range(len(df)):[cite: 1]
            discussion_val = st.session_state.get(f"dp_{idx}", df.at[idx, "Discussion Points"])[cite: 1]
            action_val = st.session_state.get(f"ap_{idx}", df.at[idx, "Action Plan"])[cite: 1]
            date_val = st.session_state.get(f"date_{idx}", df.at[idx, "Indicative Delivery Date"])[cite: 1]
            pic_val = st.session_state.get(f"pic_{idx}", df.at[idx, "Person-in-charge"])[cite: 1]
            
            rows_data.append({[cite: 1]
                "Discussion Points": discussion_val,[cite: 1]
                "Action Plan": action_val,[cite: 1]
                "Indicative Delivery Date": date_val,[cite: 1]
                "Person-in-charge": pic_val[cite: 1]
            })[cite: 1]
        
        st.session_state["df"] = pd.DataFrame(rows_data, columns=["Discussion Points", "Action Plan", "Indicative Delivery Date", "Person-in-charge"])[cite: 1]
        
        # Add Item Button
        add_col, _ = st.columns([2, 8])[cite: 1]
        with add_col:[cite: 1]
            if st.button("+ Add Item", key="add_row"):[cite: 1]
                new_row_df = pd.DataFrame([{[cite: 1]
                    "Discussion Points": "",[cite: 1]
                    "Action Plan": "",[cite: 1]
                    "Indicative Delivery Date": "",[cite: 1]
                    "Person-in-charge": ""[cite: 1]
                }])[cite: 1]
                st.session_state["df"] = pd.concat([st.session_state["df"], new_row_df], ignore_index=True)[cite: 1]
                st.rerun()[cite: 1]
        
        st.markdown('<span class="playfair-label" style="margin-top:0.75rem;">Other Discussions</span>', unsafe_allow_html=True)
        st.session_state["other_discussions"] = st.text_area(
            "Other Discussions Content",
            value=st.session_state["other_discussions"],
            height=100,
            label_visibility="collapsed"
        )

        time_range_str = f"{start_str} to {end_str}"[cite: 1]

        meeting_details = {[cite: 1]
            "date": meeting_date.strftime("%B %d, %Y"),[cite: 1]
            "time_range": time_range_str,[cite: 1]
            "location": meeting_location if meeting_location.strip() else "____________",[cite: 1]
            "company_name": client_name.strip() if client_name.strip() else "",[cite: 1]
            "prime_attendees": selected_crd,[cite: 1]
            "external_attendees": [x.strip() for x in ext_attendees_raw.split(",") if x.strip()],[cite: 1]
            "prep_name": prep_name.strip(),[cite: 1]
            "prep_desig": prep_desig.strip(),[cite: 1]
            "conf_name": conf_name.strip(),[cite: 1]
            "conf_desig": conf_desig.strip()[cite: 1]
        }[cite: 1]

        # Dual Export Section (Word DOCX and PDF)
        exp_col1, exp_col2 = st.columns(2)[cite: 1]
        
        with exp_col1:[cite: 1]
            doc_bio = export_to_word([cite: 1]
                st.session_state["df"],[cite: 1]
                meeting_details,[cite: 1]
                st.session_state["other_discussions"][cite: 1]
            )[cite: 1]
            st.download_button([cite: 1]
                label="Download Word Document (.docx)",[cite: 1]
                data=doc_bio,[cite: 1]
                file_name=f"MOM_{client_name.replace(' ', '_') if client_name else 'Report'}.docx",[cite: 1]
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",[cite: 1]
                key="btn_download_docx"[cite: 1]
            )[cite: 1]

        with exp_col2:[cite: 1]
            pdf_bio = export_to_pdf([cite: 1]
                st.session_state["df"],[cite: 1]
                meeting_details,[cite: 1]
                st.session_state["other_discussions"][cite: 1]
            )[cite: 1]
            st.download_button([cite: 1]
                label="Download PDF Document (.pdf)",[cite: 1]
                data=pdf_bio,[cite: 1]
                file_name=f"MOM_{client_name.replace(' ', '_') if client_name else 'Report'}.pdf",[cite: 1]
                mime="application/pdf",[cite: 1]
                key="btn_download_pdf"[cite: 1]
            )[cite: 1]
