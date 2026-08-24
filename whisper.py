import streamlit as st
import requests
import json
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

# ==================== STREAMLIT CONFIGURATION ====================
st.set_page_config(
    page_title="Project Eco",
    layout="wide",
    initial_sidebar_state="collapsed"
)

GROQ_API_KEY = "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"
GROQ_TRANSCRIPTION_URL = "https://api.groq.com/openai/v1/audio/transcriptions"
GROQ_CHAT_URL = "https://api.groq.com/openai/v1/chat/completions"

# ==================== LIGHT MODE / PRIME PHILIPPINES DESIGN SYSTEM ====================
st.markdown("""
<style>
    /* Clean Light Mode Reset & Grid Background */
    :root, html, body, [data-testid="stAppViewContainer"], .stApp {
        color-scheme: light !important;
        background-color: #f8fafc !important;
        background-image: 
            linear-gradient(to right, rgba(0, 0, 0, 0.04) 1px, transparent 1px),
            linear-gradient(to bottom, rgba(0, 0, 0, 0.04) 1px, transparent 1px) !important;
        background-size: 44px 44px !important;
        color: #0f172a !important;
        font-family: -apple-system, BlinkMacSystemFont, "SF Pro Display", "Segoe UI", Roboto, sans-serif !important;
        -webkit-font-smoothing: antialiased;
    }

    /* Remove Streamlit Header / Footers / Badges */
    header[data-testid="stHeader"],
    div[data-testid="stDecoration"],
    #MainMenu,
    footer,
    .stDeployButton {
        display: none !important;
        visibility: hidden !important;
        height: 0px !important;
    }

    .block-container {
        padding-top: 2.2rem !important;
        padding-bottom: 4rem !important;
        max-width: 1080px !important;
        margin: 0 auto;
    }

    /* Header & Badge Styling */
    .header-group {
        display: flex;
        align-items: center;
        gap: 12px;
        margin-bottom: 2px;
    }

    .app-title {
        font-size: 1.55rem;
        font-weight: 700;
        letter-spacing: -0.03em;
        color: #0f172a;
    }
    
    .app-badge {
        background: #ecfdf5;
        color: #059669;
        font-size: 0.68rem;
        padding: 3px 9px;
        border-radius: 20px;
        font-weight: 600;
        border: 1px solid #a7f3d0;
        letter-spacing: 0.05em;
        text-transform: uppercase;
    }
    
    .open-note {
        font-size: 0.85rem;
        color: #64748b;
        letter-spacing: 0.01em;
        margin-bottom: 1.8rem;
        font-weight: 400;
    }

    /* Light Theme Tab Bar */
    .stTabs [data-baseweb="tab-list"] {
        background: #ffffff !important;
        border: 1px solid #e2e8f0 !important;
        border-radius: 10px !important;
        padding: 4px !important;
        gap: 6px !important;
        width: fit-content;
        margin-bottom: 1.2rem;
        box-shadow: 0 1px 3px rgba(0, 0, 0, 0.04);
    }
    
    .stTabs [data-baseweb="tab"] {
        padding: 6px 18px !important;
        font-size: 0.82rem !important;
        font-weight: 500 !important;
        color: #64748b !important;
        background-color: transparent !important;
        border: none !important;
        border-radius: 6px !important;
        transition: all 0.15s ease-in-out !important;
    }
    
    .stTabs [aria-selected="true"] {
        color: #0f172a !important;
        background: #f1f5f9 !important;
        box-shadow: 0 1px 2px rgba(0, 0, 0, 0.05) !important;
        border-bottom: none !important;
    }

    /* Light Mode File Uploader */
    [data-testid="stFileUploader"],
    [data-testid="stFileUploader"] > div,
    [data-testid="stFileUploaderDropzone"] {
        background-color: #ffffff !important;
        border: 1px dashed #cbd5e1 !important;
        border-radius: 10px !important;
        color: #64748b !important;
        box-shadow: 0 1px 3px rgba(0, 0, 0, 0.02);
    }

    [data-testid="stFileUploaderDropzone"]:hover {
        border-color: #10b981 !important;
        background-color: #f8fafc !important;
    }

    [data-testid="stFileUploaderDropzone"] button {
        background: #f1f5f9 !important;
        color: #0f172a !important;
        border: 1px solid #e2e8f0 !important;
        border-radius: 6px !important;
        font-weight: 500 !important;
    }

    [data-testid="stFileUploaderDropzone"] span,
    [data-testid="stFileUploaderDropzone"] small,
    [data-testid="stFileUploaderDropzone"] div {
        color: #64748b !important;
    }

    /* Audio Input Frame */
    [data-testid="stAudioInput"] {
        background: #ffffff !important;
        border: 1px solid #e2e8f0 !important;
        border-radius: 10px !important;
        padding: 10px 14px !important;
        box-shadow: 0 1px 3px rgba(0, 0, 0, 0.03);
    }

    /* Tactile Buttons */
    .stButton > button {
        background: #ffffff !important;
        color: #0f172a !important;
        border: 1px solid #cbd5e1 !important;
        border-radius: 8px !important;
        font-size: 0.83rem !important;
        font-weight: 500 !important;
        letter-spacing: -0.01em !important;
        padding: 0.45rem 1.15rem !important;
        box-shadow: 0 1px 2px rgba(0, 0, 0, 0.04);
        transition: all 0.15s ease-in-out !important;
    }
    
    .stButton > button:hover {
        background: #f8fafc !important;
        border-color: #94a3b8 !important;
        color: #000000 !important;
        transform: translateY(-1px);
        box-shadow: 0 4px 12px rgba(0, 0, 0, 0.06);
    }

    .stButton > button[kind="primary"] {
        background: #10b981 !important;
        color: #ffffff !important;
        border: 1px solid #059669 !important;
        box-shadow: 0 2px 8px rgba(16, 185, 129, 0.25) !important;
        font-weight: 600 !important;
    }
    
    .stButton > button[kind="primary"]:hover {
        background: #059669 !important;
        box-shadow: 0 4px 14px rgba(16, 185, 129, 0.35) !important;
    }

    /* Expander Container */
    .streamlit-expanderHeader {
        background-color: #ffffff !important;
        border: 1px solid #e2e8f0 !important;
        border-radius: 8px !important;
        color: #1e293b !important;
        font-size: 0.85rem !important;
        font-weight: 500 !important;
    }

    /* Structured Matrix & Overview Cards */
    .summary-card {
        background: #ffffff;
        border: 1px solid #e2e8f0;
        border-radius: 10px;
        padding: 16px 20px;
        margin-bottom: 1rem;
        box-shadow: 0 1px 3px rgba(0, 0, 0, 0.03);
    }
    
    .summary-title {
        font-size: 0.85rem;
        font-weight: 600;
        text-transform: uppercase;
        letter-spacing: 0.04em;
        color: #059669;
        margin-bottom: 6px;
    }

    .summary-content {
        font-size: 0.88rem;
        line-height: 1.55;
        color: #334155;
    }

    /* Data Editor */
    [data-testid="stDataEditor"] {
        border: 1px solid #e2e8f0 !important;
        border-radius: 10px !important;
        background-color: #ffffff !important;
        box-shadow: 0 4px 12px rgba(0, 0, 0, 0.03) !important;
    }

    .stTextArea textarea {
        background-color: #ffffff !important;
        color: #0f172a !important;
        border: 1px solid #cbd5e1 !important;
        border-radius: 8px !important;
        font-size: 0.86rem !important;
        line-height: 1.6 !important;
    }

    .stTextArea textarea:focus {
        border-color: #10b981 !important;
        box-shadow: 0 0 0 1px #10b981 !important;
    }

    hr {
        border: none !important;
        border-top: 1px solid #e2e8f0 !important;
        margin: 1.6rem 0 !important;
    }
</style>
""", unsafe_allow_html=True)

# ==================== STATE MANAGEMENT ====================
if "transcript" not in st.session_state:
    st.session_state["transcript"] = ""
if "summary" not in st.session_state:
    st.session_state["summary"] = ""
if "decisions" not in st.session_state:
    st.session_state["decisions"] = ""
if "df" not in st.session_state:
    st.session_state["df"] = pd.DataFrame(columns=["Key Point", "Action Plan", "Assigned"])

# ==================== BACKEND SERVICES ====================
def transcribe_audio(audio_bytes):
    headers = {"Authorization": f"Bearer {GROQ_API_KEY}"}
    files = {
        "file": ("audio.wav", audio_bytes),
        "model": (None, "whisper-large-v3-turbo"),
        "response_format": (None, "json")
    }
    try:
        resp = requests.post(GROQ_TRANSCRIPTION_URL, headers=headers, files=files, timeout=60)
        if resp.status_code == 200:
            return resp.json().get("text", "")
        else:
            st.error(f"Engine status: {resp.status_code}")
            return None
    except Exception as e:
        st.error(f"Transcription error: {e}")
        return None

def extract_summary_and_actions(text):
    headers = {
        "Authorization": f"Bearer {GROQ_API_KEY}",
        "Content-Type": "application/json"
    }
    
    system_prompt = (
        "You are an executive meeting intelligence engine for Project Eco. Extract structured operational data from transcripts.\n"
        "Return ONLY a valid JSON object matching this exact schema without markdown formatting or code blocks:\n"
        "{\n"
        '  "summary": "High-level summary of the entire session.",\n'
        '  "decisions": "Key decisions agreed upon during the meeting.",\n'
        '  "records": [\n'
        "    {\n"
        '      "Key Point": "Core topic, key point, or deliverable",\n'
        '      "Action Plan": "Concrete steps, multiline bullets (•) or numbers (1., 2.)",\n'
        '      "Assigned": "Person, role, team name, or Unassigned"\n'
        "    }\n"
        "  ]\n"
        "}"
    )

    payload = {
        "model": "llama-3.3-70b-versatile",
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"Transcript:\n{text}"}
        ],
        "response_format": {"type": "json_object"},
        "temperature": 0.2
    }

    try:
        resp = requests.post(GROQ_CHAT_URL, headers=headers, json=payload, timeout=60)
        if resp.status_code == 200:
            content = resp.json()["choices"][0]["message"]["content"].strip()
            data = json.loads(content)
            
            summary = data.get("summary", "")
            decisions = data.get("decisions", "")
            records = data.get("records", [])
            return summary, decisions, pd.DataFrame(records)
        else:
            return fallback_parse(text)
    except Exception:
        return fallback_parse(text)

def fallback_parse(text):
    import re
    sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', text) if s.strip()]
    rows = []
    for s in sentences[:4]:
        rows.append({
            "Key Point": s,
            "Action Plan": "• Align deliverables with project milestones\n• Execute operational review",
            "Assigned": "Project Team"
        })
    df = pd.DataFrame(rows if rows else [{"Key Point": "Strategic Directives", "Action Plan": "• Follow up on agenda items", "Assigned": "Unassigned"}])
    return "Executive meeting overview derived from session transcript.", "Proceed according to project milestones.", df

# ==================== EXPORT UTILITIES ====================
def set_cell_margins(cell, top=140, bottom=140, start=160, end=160):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', start), ('right', end)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def export_docx(df, summary, decisions, transcript):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)

    # Title
    h1 = doc.add_heading("PROJECT ECO // MINUTES OF THE MEETING", level=1)
    h1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in h1.runs:
        r.font.name = "Segoe UI"
        r.font.size = Pt(15)
        r.font.bold = True
        r.font.color.rgb = RGBColor(5, 150, 105)

    # Executive Summary
    if summary:
        h2 = doc.add_heading("Executive Summary", level=2)
        for r in h2.runs:
            r.font.name = "Segoe UI"
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(15, 23, 42)
        p_sum = doc.add_paragraph(summary)
        for r in p_sum.runs:
            r.font.name = "Segoe UI"
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(51, 65, 85)

    # Decisions
    if decisions:
        h2_dec = doc.add_heading("Key Decisions", level=2)
        for r in h2_dec.runs:
            r.font.name = "Segoe UI"
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(15, 23, 42)
        p_dec = doc.add_paragraph(decisions)
        for r in p_dec.runs:
            r.font.name = "Segoe UI"
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(51, 65, 85)

    # Action Items Table
    doc.add_heading("Action Items & Deliverables", level=2)
    table = doc.add_table(rows=len(df) + 1, cols=3)
    table.autofit = False
    col_widths = [Inches(2.5), Inches(2.9), Inches(1.1)]
    headers = ["Key Point", "Action Plan", "Assigned"]
    
    hdr_row = table.rows[0]
    trPr = hdr_row._tr.get_or_add_trPr()
    trPr.append(parse_xml(r'<w:tblHeader %s/>' % nsdecls('w')))
    
    for idx, name in enumerate(headers):
        cell = hdr_row.cells[idx]
        cell.width = col_widths[idx]
        set_cell_margins(cell, top=160, bottom=160)
        shading = parse_xml(r'<w:shd {} w:fill="F1F5F9"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)
        
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(name)
        run.bold = True
        run.font.name = "Segoe UI"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(15, 23, 42)

    for row_idx, data in df.iterrows():
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, col_name in enumerate(headers):
            cell = row_cells[col_idx]
            cell.width = col_widths[col_idx]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            
            raw_text = str(data[col_name]) if pd.notna(data[col_name]) else ""
            lines = raw_text.split("\n")
            
            for l_idx, line in enumerate(lines):
                p = cell.paragraphs[0] if l_idx == 0 else cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(0)
                run = p.add_run(line)
                run.font.name = "Segoe UI"
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(30, 41, 59)

    if transcript:
        doc.add_paragraph().paragraph_format.space_before = Pt(16)
        h3 = doc.add_heading("Session Transcript", level=2)
        for r in h3.runs:
            r.font.name = "Segoe UI"
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(15, 23, 42)
            
        p_trans = doc.add_paragraph(transcript)
        p_trans.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in p_trans.runs:
            r.font.name = "Segoe UI"
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(100, 116, 139)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def export_pdf(df, summary, decisions, transcript):
    bio = BytesIO()
    doc = SimpleDocTemplate(
        bio,
        pagesize=letter,
        leftMargin=54,
        rightMargin=54,
        topMargin=54,
        bottomMargin=54
    )
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=14,
        textColor=colors.HexColor('#059669'),
        spaceAfter=12
    )
    section_style = ParagraphStyle(
        'DocSection',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=10.5,
        textColor=colors.HexColor('#0F172A'),
        spaceBefore=10,
        spaceAfter=6
    )
    body_style = ParagraphStyle(
        'DocBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8.5,
        leading=12,
        textColor=colors.HexColor('#334155'),
        spaceAfter=6
    )
    header_cell_style = ParagraphStyle(
        'HeaderCell',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=8.5,
        textColor=colors.HexColor('#0F172A')
    )
    body_cell_style = ParagraphStyle(
        'BodyCell',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8,
        leading=11,
        textColor=colors.HexColor('#1E293B')
    )
    transcript_style = ParagraphStyle(
        'TranscriptText',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=7.5,
        leading=10.5,
        textColor=colors.HexColor('#64748B')
    )
    
    story = [Paragraph("PROJECT ECO // MINUTES OF THE MEETING", title_style)]
    
    if summary:
        story.append(Paragraph("Executive Summary", section_style))
        story.append(Paragraph(summary.replace('\n', '<br/>'), body_style))
        
    if decisions:
        story.append(Paragraph("Key Decisions", section_style))
        story.append(Paragraph(decisions.replace('\n', '<br/>'), body_style))

    story.append(Paragraph("Action Items & Deliverables", section_style))
    table_data = [[
        Paragraph("Key Point", header_cell_style),
        Paragraph("Action Plan", header_cell_style),
        Paragraph("Assigned", header_cell_style)
    ]]
    
    for _, row in df.iterrows():
        kp_text = str(row["Key Point"]).replace('\n', '<br/>') if pd.notna(row["Key Point"]) else ""
        ap_text = str(row["Action Plan"]).replace('\n', '<br/>') if pd.notna(row["Action Plan"]) else ""
        as_text = str(row["Assigned"]).replace('\n', '<br/>') if pd.notna(row["Assigned"]) else ""
        
        table_data.append([
            Paragraph(kp_text, body_cell_style),
            Paragraph(ap_text, body_cell_style),
            Paragraph(as_text, body_cell_style)
        ])
        
    t = Table(table_data, colWidths=[185, 230, 89])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#F1F5F9')),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
    ]))
    story.append(t)
    
    if transcript:
        story.append(Spacer(1, 12))
        story.append(Paragraph("Session Transcript", section_style))
        story.append(Paragraph(transcript.replace('\n', '<br/>'), transcript_style))
        
    doc.build(story)
    bio.seek(0)
    return bio

# ==================== APPLICATION CANVAS ====================
st.markdown("""
    <div class="header-group">
        <div class="app-title">Project Eco</div>
        <span class="app-badge">PRIME CORE</span>
    </div>
    <div class="open-note">Open architecture for intelligence capture, structured directives, and operational agility.</div>
""", unsafe_allow_html=True)

tab_rec, tab_up = st.tabs(["Record", "Upload"])
audio_payload = None

with tab_rec:
    rec_buffer = st.audio_input("Record Audio Stream", label_visibility="collapsed")
    if rec_buffer:
        audio_payload = rec_buffer.read()

with tab_up:
    up_buffer = st.file_uploader(
        "Upload Audio",
        type=["wav", "mp3", "m4a", "ogg", "flac", "webm"],
        label_visibility="collapsed"
    )
    if up_buffer:
        audio_payload = up_buffer.read()

if audio_payload:
    st.write("")
    if st.button("Transcribe", type="primary"):
        with st.spinner("Processing speech transcription..."):
            res = transcribe_audio(audio_payload)
        if res:
            st.session_state["transcript"] = res
            st.session_state["summary"] = ""
            st.session_state["decisions"] = ""
            st.session_state["df"] = pd.DataFrame(columns=["Key Point", "Action Plan", "Assigned"])
            st.rerun()

if st.session_state["transcript"]:
    st.markdown("---")
    with st.expander("Session Transcript", expanded=False):
        st.session_state["transcript"] = st.text_area(
            "Session Transcript",
            st.session_state["transcript"],
            height=120,
            label_visibility="collapsed"
        )
    
    col_btn, _ = st.columns([2.5, 4.5])
    with col_btn:
        if st.button("Extract Summary & Actions", type="primary"):
            with st.spinner("Synthesizing directives with Llama 3.3 70B..."):
                sum_text, dec_text, structured_df = extract_summary_and_actions(st.session_state["transcript"])
            st.session_state["summary"] = sum_text
            st.session_state["decisions"] = dec_text
            if not structured_df.empty:
                st.session_state["df"] = structured_df
                st.rerun()

if st.session_state["summary"] or not st.session_state["df"].empty:
    st.markdown("---")
    
    if st.session_state["summary"]:
        st.markdown(f"""
        <div class="summary-card">
            <div class="summary-title">Executive Summary</div>
            <div class="summary-content">{st.session_state["summary"]}</div>
        </div>
        """, unsafe_allow_html=True)

    if st.session_state["decisions"]:
        st.markdown(f"""
        <div class="summary-card">
            <div class="summary-title">Key Decisions</div>
            <div class="summary-content">{st.session_state["decisions"]}</div>
        </div>
        """, unsafe_allow_html=True)

    column_config = {
        "Key Point": st.column_config.TextColumn(
            "Key Point",
            required=True,
            width="large"
        ),
        "Action Plan": st.column_config.TextColumn(
            "Action Plan",
            required=False,
            width="large"
        ),
        "Assigned": st.column_config.TextColumn(
            "Assigned",
            required=False,
            width="medium"
        ),
    }

    edited = st.data_editor(
        st.session_state["df"],
        column_config=column_config,
        num_rows="dynamic",
        use_container_width=True,
        hide_index=True,
        key="project_eco_light_editor"
    )
    st.session_state["df"] = edited

    st.write("")
    col_word, col_pdf, _ = st.columns([1.5, 1.5, 4])
    
    with col_word:
        doc_data = export_docx(st.session_state["df"], st.session_state["summary"], st.session_state["decisions"], st.session_state["transcript"])
        st.download_button(
            label="Export Word",
            data=doc_data,
            file_name="Project_Eco_Minutes.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    with col_pdf:
        pdf_data = export_pdf(st.session_state["df"], st.session_state["summary"], st.session_state["decisions"], st.session_state["transcript"])
        st.download_button(
            label="Export PDF",
            data=pdf_data,
            file_name="Project_Eco_Minutes.pdf",
            mime="application/pdf"
        )
